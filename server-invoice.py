from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from num2words import num2words
import os
import smtplib
from email.mime.text import MIMEText
import html
from PIL import Image, ImageDraw, ImageFont
from datetime import datetime

app = FastAPI()

# Add CORS middleware to allow all origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allow all methods (GET, POST, etc.)
    allow_headers=["*"],  # Allow all headers
)

# Pydantic models for request validation
class Item(BaseModel):
    description: str
    sacCode: str
    amount: str  # Kept as str to match input, will parse to float

class InvoiceRequest(BaseModel):
    date: str
    invoiceNo: str
    billTo: str
    instamojoPaymentId: str
    items: list[Item]
    email: str
    fees: str  # Field for Instamojo fees (kept for backward compatibility, not used)

class UpgradeInvoiceRequest(BaseModel):
    billTo: str
    email: str
    instamojoPaymentId: str
    packageName: str
    amount: str

@app.post("/api/generate-invoice")
async def generate_invoice(request: InvoiceRequest):
    try:
        # Extract data from request
        date = request.date
        invoice_no = request.invoiceNo
        bill_to = request.billTo
        instamojo_payment_id = request.instamojoPaymentId
        items = request.items
        recipient_email = request.email

        # Calculate totals
        basic_value = sum(float(item.amount) for item in items)
        taxable_value = basic_value
        cgst_rate = 0.09  # 9% CGST
        sgst_rate = 0.09  # 9% SGST
        igst_rate = 0.0   # 0% IGST
        cgst = taxable_value * cgst_rate
        sgst = taxable_value * sgst_rate
        igst = taxable_value * igst_rate
        total_tax = cgst + sgst + igst
        subtotal = taxable_value
        total = taxable_value + total_tax

        # Calculate Instamojo fees: 3% of total + 18% GST on that 3% + Rs. 3 fixed
        base_fee = total * 0.03  # 3% of total
        fixed_fee = 3.0  # Rs. 3 fixed
        instamojo_taxable_value = base_fee + fixed_fee
        gst_on_fee = instamojo_taxable_value * 0.18  # 18% GST on the 3% fee
        instamojo_fees = instamojo_taxable_value + gst_on_fee

        # Calculate grand total
        grand_total = total + instamojo_fees
        amount_in_words = num2words(grand_total, lang='en').replace(',', '').title() + " Rupees Only"

        # Create DOCX document
        doc = Document()

        # Header
        doc.add_heading("Karnataka Chapter of The India Society for Assisted Reproduction", level=2).alignment = 1  # Center
        doc.add_paragraph("No.1, 1st Floor, UMA Admiralty,\nBannerugatta Road, Bangalore - 560029.\nGSTIN: 29AABAK4261H2ZL", style='Normal').alignment = 1

        # Invoice Details
        doc.add_heading("GST INVOICE", level=3)
        doc.add_paragraph(f"DATE: {date}")
        doc.add_paragraph(f"INVOICE NO: {invoice_no}")
        doc.add_paragraph("10TH ANNUAL CONFERENCE KISAR-2025")

        # Bill To
        doc.add_paragraph("BILL TO", style='Normal').runs[0].bold = True
        doc.add_paragraph(bill_to)

        # Instamojo Payment ID
        doc.add_paragraph(f"Instamojo Payment ID: {instamojo_payment_id}")

        # Items Table
        items_table = doc.add_table(rows=1, cols=3)
        items_table.style = 'Table Grid'
        hdr_cells = items_table.rows[0].cells
        hdr_cells[0].text = "DESCRIPTION"
        hdr_cells[1].text = "SAC Code"
        hdr_cells[2].text = "Amount in Rs."
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True

        for item in items:
            row_cells = items_table.add_row().cells
            row_cells[0].text = item.description
            row_cells[1].text = item.sacCode
            row_cells[2].text = f"{float(item.amount):.2f}"

        # Spacer paragraph before totals table
        doc.add_paragraph("", style='Normal').paragraph_format.space_after = 240  # Approx 12pt spacing

        # Totals Table (Includes all requested fields)
        totals_table = doc.add_table(rows=10, cols=2)
        totals_table.style = 'Table Grid'
        totals_table.autofit = False
        totals_table.columns[0].width = 3000000  # Approx 50% width
        totals_table.columns[1].width = 3000000
        totals_table.alignment = 2  # Right align

        totals_data = [
            ("Basic Value:", f"₹{basic_value:.2f}"),
            ("Taxable Value:", f"₹{taxable_value:.2f}"),
            ("CGST (9%):", f"₹{cgst:.2f}"),
            ("SGST (9%):", f"₹{sgst:.2f}"),
            ("IGST (0%):", f"₹{igst:.2f}"),
            ("Total Tax:", f"₹{total_tax:.2f}"),
            ("Subtotal:", f"₹{subtotal:.2f}"),
            ("Total:", f"₹{total:.2f}"),
            ("Payment convenience Fee (3% on total amount + Rs.3 + 18% GST ):", f"₹{instamojo_fees:.2f}"),
            ("Grand Total:", f"₹{grand_total:.2f}"),
        ]

        for i, (label, value) in enumerate(totals_data):
            row_cells = totals_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            row_cells[1].paragraphs[0].alignment = 2  # Right align value
            if label in ["Total:", "Grand Total:"]:
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[1].paragraphs[0].runs[0].bold = True

        # Spacer paragraph before Amount in Words
        doc.add_paragraph("", style='Normal').paragraph_format.space_after = 480  # Approx 24pt spacing

        # Amount in Words
        doc.add_paragraph(f"Amount in Words: {amount_in_words}")

        # Remittance Section
        doc.add_paragraph("Please Remit the Amount to Below Account Number", style='Normal').runs[0].bold = True
        doc.add_paragraph("Saving A/c No.: 029901005524")
        doc.add_paragraph("ICICI Bank, Jayanagar 9th Block Bangalore")
        doc.add_paragraph("IFSC Code: ICICI0000299")
        doc.add_paragraph("PAN: AABAK4261H")

        # Save the document
        sanitized_bill_to = "".join(c if c.isalnum() else "_" for c in bill_to)
        filename = f"{sanitized_bill_to}_{instamojo_payment_id}.docx"
        folder_path = os.path.join(os.getcwd(), "invoices", "new")
        os.makedirs(folder_path, exist_ok=True)
        file_path = os.path.join(folder_path, filename)

        doc.save(file_path)

        # Send confirmation email with updated HTML/CSS template
        sender_email = "labslyxn@gmail.com"
        sender_password = "iwjnveuscunamwgs"
        cc_email = "mfaisal.pla@gmail.com"

        # Generate items table for email
        items_html = "".join(
            f"""
            <tr>
                <td style="padding: 10px; border-bottom: 1px solid #eee;">{item.description}</td>
                <td style="padding: 10px; border-bottom: 1px solid #eee; text-align: right;">₹{float(item.amount):.2f}</td>
            </tr>
            """ for item in items
        )

        # Email template with all fields
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    background-color: #f4f4f4;
                    margin: 0;
                    padding: 0;
                }}
                .container {{
                    max-width: 600px;
                    margin: 20px auto;
                    background-color: #ffffff;
                    border-radius: 10px;
                    box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                    overflow: hidden;
                }}
                .header {{
                    background-color: #4CAF50;
                    color: white;
                    padding: 20px;
                    text-align: center;
                    border-top-left-radius: 10px;
                    border-top-right-radius: 10px;
                }}
                .content {{
                    padding: 20px;
                }}
                .footer {{
                    background-color: #f4f4f4;
                    padding: 10px;
                    text-align: center;
                    font-size: 12px;
                    color: #666;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}
                th, td {{
                    padding: 10px;
                    text-align: left;
                }}
                th {{
                    background-color: #f0f0f0;
                    font-weight: bold;
                }}
                .total {{
                    font-weight: bold;
                    font-size: 16px;
                }}
                .payment-id {{
                    font-size: 14px;
                    color: #333;
                    margin-top: 10px;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h2>Thank You for Registering!</h2>
                </div>
                <div class="content">
                    <p>Dear {bill_to},</p>
                    <p>Thank you for registering for the <strong>10th Annual Conference 2025 - KISAR</strong>. Below are your registration details:</p>
                    <table>
                        <thead>
                            <tr>
                                <th>Item</th>
                                <th>Amount</th>
                            </tr>
                        </thead>
                        <tbody>
                            {items_html}
                        </tbody>
                    </table>
                    <table>
                        <tr>
                            <td style="padding: 10px;">Subtotal:</td>
                            <td style="padding: 10px; text-align: right;">₹{subtotal:.2f}</td>
                        </tr>
                        <tr>
                            <td style="padding: 10px;">CGST (9%):</td>
                            <td style="padding: 10px; text-align: right;">₹{cgst:.2f}</td>
                        </tr>
                        <tr>
                            <td style="padding: 10px;">SGST (9%):</td>
                            <td style="padding: 10px; text-align: right;">₹{sgst:.2f}</td>
                        </tr>
                        <tr>
                            <td style="padding: 10px;">Total:</td>
                            <td style="padding: 10px; text-align: right; font-weight: bold;">₹{total:.2f}</td>
                        </tr>
                        <tr>
                            <td style="padding: 10px;">Payment convenience Fee (3% on total amount + Rs.3 + 18% GST ):</td>
                            <td style="padding: 10px; text-align: right;">₹{instamojo_fees:.2f}</td>
                        </tr>
                        <tr>
                            <td style="padding: 10px;">Grand Total:</td>
                            <td style="padding: 10px; text-align: right; font-weight: bold;">₹{grand_total:.2f}</td>
                        </tr>
                    </table>
                    <p class="payment-id">Instamojo Payment ID: {instamojo_payment_id}</p>
                    <p>We look forward to seeing you at the event!</p>
                    <p>Best regards,<br>Karnataka Chapter of ISAR</p>
                </div>
                <div class="footer">
                    <p>Contact us at: kisar.office@gmail.com | © 2025 KISAR</p>
                    <p>Powered By: © LYXN LABS</p>
                </div>
            </div>
        </body>
        </html>
        """

        msg = MIMEText(html_content, 'html')
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Cc'] = cc_email
        msg['Subject'] = "Registration Confirmation - 10th Annual Conference 2025 KISAR -Updated Invoice"

        recipients = [recipient_email, cc_email]

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipients, msg.as_string())

        return {"status": "OK"}

    except Exception as e:
        print(f"Error generating invoice: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate invoice")

@app.post("/api/generate-upgrade-invoice")
async def generate_upgrade_invoice(request: UpgradeInvoiceRequest):
    try:
        # Extract data from request
        bill_to = request.billTo
        recipient_email = request.email
        instamojo_payment_id = request.instamojoPaymentId
        package_name = request.packageName
        basic_value = float(request.amount)  # Use request.amount as basic_value

        # Calculate totals (no GST applied for upgrades)
        total = basic_value

        # Calculate Instamojo fees: 3% of total + 18% GST on that 3% + Rs. 3 fixed
        base_fee = total * 0.03  # 3% of total
        fixed_fee = 3.0  # Rs. 3 fixed
        instamojo_taxable_value = base_fee + fixed_fee
        gst_on_fee = instamojo_taxable_value * 0.18  # 18% GST on the 3% fee
        instamojo_fees = instamojo_taxable_value + gst_on_fee

        # Calculate grand total
        grand_total = total + instamojo_fees

        # Send confirmation email with updated HTML/CSS template
        sender_email = "labslyxn@gmail.com"
        sender_password = "iwjnveuscunamwgs"
        cc_email = "mfaisal.pla@gmail.com"

        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    background-color: #f4f4f4;
                    margin: 0;
                    padding: 0;
                }}
                .container {{
                    max-width: 600px;
                    margin: 20px auto;
                    background-color: #ffffff;
                    border-radius: 10px;
                    box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                    overflow: hidden;
                }}
                .header {{
                    background-color: #4CAF50;
                    color: white;
                    padding: 20px;
                    text-align: center;
                    border-top-left-radius: 10px;
                    border-top-right-radius: 10px;
                }}
                .content {{
                    padding: 20px;
                }}
                .footer {{
                    background-color: #f4f4f4;
                    padding: 10px;
                    text-align: center;
                    font-size: 12px;
                    color: #666;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}
                th, td {{
                    padding: 10px;
                    text-align: left;
                }}
                th {{
                    background-color: #f0f0f0;
                    font-weight: bold;
                }}
                .total {{
                    font-weight: bold;
                    font-size: 16px;
                }}
                .payment-id {{
                    font-size: 14px;
                    color: #333;
                    margin-top: 10px;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h2>Package Upgrade Confirmation</h2>
                </div>
                <div class="content">
                    <p>Dear {bill_to},</p>
                    <p>Your package for the <strong>10th Annual Conference 2025 - KISAR</strong> has been successfully upgraded. Below are the details:</p>
                    <table>
                        <thead>
                            <tr>
                                <th>Package</th>
                                <th>Amount</th>
                            </tr>
                        </thead>
                        <tbody>
                            <tr>
                                <td style="padding: 10px; border-bottom: 1px solid #eee;">{package_name}</td>
                                <td style="padding: 10px; border-bottom: 1px solid #eee; text-align: right;">₹{basic_value:.2f}</td>
                            </tr>
                        </tbody>
                    </table>
                    <table>
                        <tr>
                            <td style="padding: 10px;">Subtotal:</td>
                            <td style="padding: 10px; text-align: right;">₹{total:.2f}</td>
                        </tr>
                        <tr>
                            <td style="padding: 10px;">Payment convenience Fee (3% on total amount +  Rs.3 +18% GST ):</td>
                            <td style="padding: 10px; text-align: right;">₹{instamojo_fees:.2f}</td>
                        </tr>
                        <tr>
                            <td style="padding: 10px;">Grand Total:</td>
                            <td style="padding: 10px; text-align: right; font-weight: bold;">₹{grand_total:.2f}</td>
                        </tr>
                    </table>
                    <p class="payment-id">Instamojo Payment ID: {instamojo_payment_id}</p>
                    <p>We look forward to seeing you at the event!</p>
                    <p>Best regards,<br>Karnataka Chapter of ISAR</p>
                </div>
                <div class="footer">
                    <p>Contact us at: kisar.office@gmail.com | © 2025 KISAR</p>
                    <p>Powered By: © LYXN LABS</p>
                </div>
            </div>
        </body>
        </html>
        """

        msg = MIMEText(html_content, 'html')
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Cc'] = cc_email
        msg['Subject'] = "Package Upgrade Confirmation - 10th Annual Conference 2025 KISAR"

        recipients = [recipient_email, cc_email]

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipients, msg.as_string())

        return {"status": "OK"}

    except Exception as e:
        print(f"Error generating upgrade invoice: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate upgrade invoice")

class CertificateData(BaseModel):
    role: str
    fullName: str
    medicalCouncilNumber: str
    medicalCouncilState: str
    city: str

@app.post("/api/generate-certificate")
async def generate_certificate(data: CertificateData):
    try:
        # Select the template based on role
        template_path = "template/faculty.png" if data.role == "Faculty" else "template/delegate.png"

        # Verify template exists
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail=f"Template {template_path} not found")

        # Open the image
        image = Image.open(template_path)
        draw = ImageDraw.Draw(image)

        # Use a default font (Pillow doesn't bundle fonts, so use a system font or provide a path to a .ttf file)
        try:
            font = ImageFont.truetype("apercumovistarbold.ttf", 50)
            smaller_font = ImageFont.truetype("apercumovistarbold.ttf", 30)  # Smaller font for adjustments
        except IOError:
            # Fallback to default font if not available
            font = ImageFont.load_default()
            smaller_font = ImageFont.load_default()  # Fallback won't support size change, but included for completeness

        # Remove trailing spaces and convert all data to uppercase
        full_name_upper = data.fullName.strip().upper()
        medical_council_number_upper = data.medicalCouncilNumber.strip().upper()
        medical_council_state_upper = data.medicalCouncilState.strip().upper()

        # Convert all data to uppercase
        full_name_upper = data.fullName.strip().upper()
        medical_council_number_upper = data.medicalCouncilNumber.strip().upper()
        medical_council_state_upper = data.medicalCouncilState.strip().upper()
        city_upper = data.city.strip().upper()

        # Full Name: 100px below top, centered (unchanged as it fits perfectly)
        full_name_bbox = draw.textbbox((0, 0), full_name_upper, font=font)
        full_name_width = full_name_bbox[2] - full_name_bbox[0]
        image_width = image.width
        full_name_x = (image_width - full_name_width) // 2
        draw.text((full_name_x, 640), full_name_upper, fill="black", font=font)

        # Medical Council Number (left) and State (right) on the next line
        # Medical Council Number: Check length and split if too long
        medical_council_bbox = draw.textbbox((0, 0), medical_council_number_upper, font=font)
        medical_council_width = medical_council_bbox[2] - medical_council_bbox[0]
        max_medical_council_width = 200  # Adjust this based on available space (estimated for x=900 to fit before state)

        if medical_council_width > max_medical_council_width:
            # Use smaller font and split into two lines
            medical_council_font = smaller_font
            # Split the medical council number roughly in half
            split_index = len(medical_council_number_upper) // 2
            part1 = medical_council_number_upper[:split_index]
            part2 = medical_council_number_upper[split_index:]
            # Draw the two parts one below the other
            part1_bbox = draw.textbbox((0, 0), part1, font=medical_council_font)
            part1_width = part1_bbox[2] - part1_bbox[0]
            draw.text((900, 730), part1, fill="black", font=medical_council_font)
            draw.text((900, 770), part2, fill="black", font=medical_council_font)
        else:
            # Use original font and draw in one line
            draw.text((900, 750), medical_council_number_upper, fill="black", font=font)

        # State (right): Reduce font size for longer lengths
        state_bbox = draw.textbbox((0, 0), medical_council_state_upper, font=font)
        state_width = state_bbox[2] - state_bbox[0]
        max_state_width = 300  # Adjust this based on available space (estimated for x=1250)

        state_font = smaller_font if state_width > max_state_width else font
        # Recalculate bbox with the chosen font
        state_bbox = draw.textbbox((0, 0), medical_council_state_upper, font=state_font)
        state_width = state_bbox[2] - state_bbox[0]
        state_x = 1250
        draw.text((state_x, 750), medical_council_state_upper, fill="black", font=state_font)

        # City: Reduce font size for longer lengths and adjust x-coordinate if smaller font is used
        city_bbox = draw.textbbox((0, 0), city_upper, font=font)
        city_width = city_bbox[2] - city_bbox[0]
        max_city_width = 600  # Adjust this based on available space (estimated for centered text)

        city_font = smaller_font if city_width > max_city_width else font
        # Recalculate bbox with the chosen font
        city_bbox = draw.textbbox((0, 0), city_upper, font=city_font)
        city_width = city_bbox[2] - city_bbox[0]
        # Adjust x-coordinate to 650 if smaller font is used, otherwise keep at 700
        city_x = 650 if city_font == smaller_font else 700
        draw.text((city_x, 850), city_upper, fill="black", font=city_font)

        # Ensure the certificates directory exists
        os.makedirs("certificates", exist_ok=True)

        # Save the modified image as a PDF
        output_filename = f"certificates/{full_name_upper}.pdf"
        image.save(output_filename, "PDF", resolution=100.0)

        # Return the PDF file as a response for download
        return FileResponse(
            path=output_filename,
            filename=f"{full_name_upper}.pdf",
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={full_name_upper}.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating certificate: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=4000)